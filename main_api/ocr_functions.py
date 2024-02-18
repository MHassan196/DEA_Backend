import difflib
from django.http import JsonResponse
import numpy as np
import openpyxl
import xlrd
import docx
from PyPDF2 import PdfReader
import PyPDF2
import fitz
import pytesseract
import tabula
import pandas as pd
from pdf2image import convert_from_bytes
from pytesseract import image_to_string
from PIL import Image
import io
from io import BytesIO
import re
import tempfile
import os 
import json
 
 
# Set the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# def read_excel(uploaded_file):
#     if uploaded_file.content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
#         wb = openpyxl.load_workbook(uploaded_file)
#         columns = wb.active[1]
#         rows = []
#         for row in wb.active.iter_rows(min_row=2, values_only=True):
#             rows.append(dict(zip([col.value for col in columns], row)))
#         return {"columns": [col.value for col in columns], "data": rows}
#     elif uploaded_file.content_type == 'application/vnd.ms-excel':
#         wb = xlrd.open_workbook(file_contents=uploaded_file.read())
#         sheet = wb.sheet_by_index(0)
#         columns = sheet.row_values(0)
#         rows = []
#         for row_num in range(1, sheet.nrows):
#             row = sheet.row_values(row_num)
#             rows.append(dict(zip(columns, row)))
#         return {"columns": columns, "data": rows}
#     else:
#         raise ValueError("Unsupported file format for Excel.")

def read_excel(uploaded_file):
    if uploaded_file.content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        wb = openpyxl.load_workbook(uploaded_file)
        columns = wb.active[1]
        rows = []
        for row in wb.active.iter_rows(min_row=2, values_only=True):
            rows.append(dict(zip([col.value for col in columns], row)))
        return rows
    elif uploaded_file.content_type == 'application/vnd.ms-excel':
        wb = xlrd.open_workbook(file_contents=uploaded_file.read())
        sheet = wb.sheet_by_index(0)
        columns = sheet.row_values(0)
        rows = []
        for row_num in range(1, sheet.nrows):
            row = sheet.row_values(row_num)
            rows.append(dict(zip(columns, row)))
        return rows
    else:
        raise ValueError("Unsupported file format for Excel.")

# def ocr_image(image):
#     return pytesseract.image_to_string(image)
    
# def ocr_image(image):
#     # Open the image file
#     img = Image.open(image)

#     # Use pytesseract to do OCR on the image
#     ocr_text = pytesseract.image_to_string(img)

#     # Process the OCR text into structured data
#     ocr_data = process_ocr_data(ocr_text)

#     return ocr_data


# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract potential column names from the first line
#     columns = re.findall(r'\S+', lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line using whitespace as a separator
#         row = re.findall(r'\S+', line)
        
#         # Check if the row has at least one element
#         if row:
#             # Create a dictionary with variable keys based on the number of columns
#             data.append(dict(zip(columns, row)))

#     return {"columns": columns, "data": data} 
    
# def ocr_image(image):
#     # Open the image file
#     img = Image.open(image)

#     # Use pytesseract to do OCR on the image
#     ocr_text = pytesseract.image_to_string(img)

#     # Process the OCR text into structured data
#     ocr_data = process_ocr_data(ocr_text)

#     return ocr_data

# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract column names from the first line
#     columns = re.findall(r'\b\w+\b', lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         row = re.findall(r'\b\w+\b', line)
#         if row:
#             data.append(dict(zip(columns, row)))

#     return {"columns": columns, "data": data}
    
# import cv2
# # import pytesseract
# # import re

# def ocr_image(image_path):
#     # Read the image using OpenCV
#     img = cv2.imread(image_path)

#     # Convert the image to grayscale
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

#     # Use adaptive thresholding to emphasize the text lines
#     _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

#     # Find contours in the image
#     contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

#     # Iterate through contours and filter for potential text contours based on area
#     text_contours = [contour for contour in contours if cv2.contourArea(contour) > 50]

#     # Extract text using OCR from the identified text contours
#     text = ""
#     for contour in text_contours:
#         x, y, w, h = cv2.boundingRect(contour)
#         text_roi = gray[y:y+h, x:x+w]

#         # Use pytesseract to extract text from the text region
#         text += pytesseract.image_to_string(text_roi)

#     # Remove unwanted characters from the text
#     text = re.sub(r'[\|_]', '', text)

#     return text

# # Call the function to extract data from the image
# image_path = input("Image Path: ")
# print("\n\n")
# result = ocr_image(image_path)

# # Display the result
# print(result)

import cv2
    
# def ocr_image(uploaded_file):
#     # Open the image file
#     img = cv2.imread(np.fromstring(uploaded_file.read(), np.uint8), cv2.IMREAD_UNCHANGED)

#     # Convert the image to grayscale
#     gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

#     # Use adaptive thresholding to emphasize the table lines
#     _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

#     # Find contours in the image
#     contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

#     # Iterate through contours and filter for potential table contours based on area
#     table_contours = [contour for contour in contours if cv2.contourArea(contour) > 1000]

#     # Draw the identified table contours on a copy of the original image
#     img_with_contours = img.copy()
#     cv2.drawContours(img_with_contours, table_contours, -1, (0, 255, 0), 2)

#     # Display the image with identified contours (for visualization purposes)
#     # cv2.imshow("Image with Contours", img_with_contours)
#     # cv2.waitKey(0)
#     # cv2.destroyAllWindows()

#     # Extract text using OCR from the identified table contours
#     table_text = ""
#     for contour in table_contours:
#         x, y, w, h = cv2.boundingRect(contour)
#         table_roi = img[y:y + h, x:x + w]

#         # Use pytesseract to extract text from the table region
#         table_text += pytesseract.image_to_string(table_roi)

#     # Process the extracted table text
#     ocr_data = process_ocr_data(table_text)

#     return ocr_data

# # Add the process_ocr_data function here as well

# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract potential column names from the first line
#     columns = re.findall(r'\S+', lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line using whitespace as a separator
#         row = re.findall(r'\S+', line)

#         # Check if the row has at least one element
#         if row:
#             # Create a dictionary with variable keys based on the number of columns
#             data.append(dict(zip(columns, row)))

#     return {"columns": columns, "data": data}

import io

def ocr_image(image_data):
    try:
        # Convert the image data to a NumPy array
        nparr = np.frombuffer(image_data, np.uint8)

        # Check if the array is empty
        if nparr.size == 0:
            raise ValueError("Empty image data")

        # Decode the image using OpenCV
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        # Convert the image to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Use adaptive thresholding to emphasize the text lines
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)

        # Find contours in the image
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        # Iterate through contours and filter for potential text contours based on area
        text_contours = [contour for contour in contours if cv2.contourArea(contour) > 50]

        # Extract text using OCR from the identified text contours
        text = ""
        for contour in text_contours:
            x, y, w, h = cv2.boundingRect(contour)
            text_roi = gray[y:y+h, x:x+w]

            # Use pytesseract to extract text from the text region
            word_text = pytesseract.image_to_string(text_roi)
            text += word_text + '\n'

        # Remove unwanted characters from the text
        text = re.sub(r'[\|_]', '', text)
        print(text)

        # Parse the tabular data
        parsed_data = parse_table(text)

        return parsed_data

    except Exception as e:
        print(f"Error in ocr_image: {e}")
        return ""
    
import csv

# def parse_table(text):
#     # Split the text into rows
#     rows = text.strip().split('\n')

#     # Extract the columns (first row)
#     columns = re.split(r'\s{2,}', rows[0].strip())

#     # Extract the data (remaining rows)
#     data = []
#     for row in rows[1:]:
#         # Use regular expression to find patterns resembling values
#         values = re.findall(r'[^\s"]+|"[^"]+"', row)

#         # Remove quotes from values if present
#         values = [v.strip('"') for v in values]

#         # Create a dictionary if values are present
#         if values:
#             row_dict = dict(zip(columns, values))
#             data.append(row_dict)

#     return {'columns': columns, 'data': data}

def parse_table(text):
    # Split the text into rows
    rows = text.strip().split('\n')

    # Use regular expression to extract column names considering irregular spacing
    header_match = re.match(r'\s*(.*?)\s*$', rows[0])
    header = re.split(r'\s+', header_match.group(1).strip()) if header_match else []

    # Extract the data (remaining rows)
    # Extract the data (remaining rows)
    data = []
    for row in rows[1:]:
        # Use regular expression to find patterns resembling values
        values = re.findall(r'[^\s"]+|"[^"]+"', row)

        # Remove quotes from values if present
        values = [v.strip('"') for v in values]

        # Create a new dictionary for each row
        current_row_dict = dict(zip(header, values))

        # Check if the row has any non-empty values before adding it to the result
        if any(current_row_dict.values()):
            data.append(current_row_dict)

    return {'columns': header, 'data': data}

# def parse_table(text):
#     # Split the text into rows
#     rows = text.strip().split('\n')

#     # Use regular expression to extract column names considering irregular spacing
#     header_match = re.match(r'\s*(.*?)\s*$', rows[0])
#     header = re.split(r'\s+', header_match.group(1).strip()) if header_match else []

#     # Extract the data (remaining rows)
#     data = []
#     for row in rows[1:]:
#         # Use regular expression to find patterns resembling values
#         values = re.findall(r'[^\s"]+|"[^"]+"', row)

#         # Remove quotes from values if present
#         values = [v.strip('"') for v in values]

#         # Create a dictionary for each row with header names as keys
#         current_row_dict = {header[i]: values[i] for i in range(min(len(header), len(values)))}

#         # Check if the row has any non-empty values before adding it to the result
#         if any(current_row_dict.values()):
#             data.append(current_row_dict)

#     return data, header


# Add the process_ocr_data function here as well

def process_ocr_data(ocr_text):
    # Process the OCR text into lines
    lines = ocr_text.split('\n')

    # Extract potential column names from the first line
    columns = re.findall(r'\S+', lines[0])

    # Extract data from subsequent lines
    data = []
    for line in lines[1:]:
        # Split the line using whitespace as a separator
        row = re.findall(r'\S+', line)

        # Check if the row has at least one element
        if row:
            # Create a dictionary with variable keys based on the number of columns
            data.append(dict(zip(columns, row)))

    return {"columns": columns, "data": data}


# def ocr_image(image_path):
#     # Open the image file
#     img = Image.open(image_path)

#     # Use pytesseract to do OCR on the image
#     ocr_text = pytesseract.image_to_string(img)

#     # Process the OCR text into structured data
#     ocr_data = process_ocr_data(ocr_text)

#     return ocr_data


# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract potential column names from the header line
#     columns = extract_columns(lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line using a combination of spaces and vertical bars as separators
#         row = re.split(r'\s*\|+\s*|\s+', line.strip())

#         # Check if the row has the same number of elements as columns
#         if len(row) == len(columns):
#             # Create a dictionary with flexible keys based on the columns
#             data.append(dict(zip(columns, row)))

#     return {"columns": columns, "data": data}


# def extract_columns(header_line):
#     # Extract potential column names based on the structure of the header line
#     # Use a combination of spaces and vertical bars as separators
#     potential_columns = re.split(r'\s*\|+\s*|\s+', header_line.strip())

#     # Process potential columns to identify variations and handle complex cases
#     columns = []
#     for col in potential_columns:
#         col = re.sub(r'[^a-zA-Z0-9 ]', '', col)
#         # Check for variations in column names (e.g., "First Name" or "Last Name")
#         variations = re.split(r'\s+', col.strip())
#         columns.extend(variations)

#     # Filter out any empty strings
#     columns = [col for col in columns if col]

#     # Use vertical alignment to detect column boundaries more effectively
#     aligned_columns = []
#     for col in columns:
#         alignment_counts = [sum(1 for char in row if char == '|') for row in potential_columns]
#         max_alignment_index = alignment_counts.index(max(alignment_counts))
#         aligned_columns.append(potential_columns[max_alignment_index].strip())

#     return aligned_columns
    

# def ocr_image(image_path):
#     # Open the image file
#     img = Image.open(image_path)

#     # Use pytesseract to do OCR on the image
#     ocr_text = pytesseract.image_to_string(img)

#     # Process the OCR text into structured data
#     extracted_data = process_ocr_data(ocr_text)

#     return extracted_data

# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract potential column names from the header line
#     columns = extract_columns(lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line using a combination of spaces and vertical bars as separators
#         row = re.split(r'\s*\|\s*|\s+', line.strip())

#         # Check if the row has the same number of elements as columns
#         if len(row) == len(columns):
#             # Create a dictionary with flexible keys based on the columns
#             data_row = dict(zip(columns, row))

#             # Further process data elements, combining split values
#             for key, value in data_row.items():
#                 if key != "Address":
#                     data_row[key] = ' '.join(value)

#             data.append(data_row)

#     return {"columns": columns, "data": data}

# def extract_columns(header_line):
#     # Extract potential column names based on the structure of the header line
#     # Use a combination of spaces and vertical bars as separators
#     columns = re.split(r'\s*\|\s*|\s+', header_line.strip())

#     # Filter out any empty strings or unwanted characters
#     columns = [col for col in columns if col]

#     return columns
    

# def extract_table_from_image(image_file):

#     # Convert the uploaded image file to a NumPy array
#     image = Image.open(image_file)

#     # Use pytesseract to do OCR on the image
#     table_text = pytesseract.image_to_string(image)

#     # Process the OCR text into structured data
#     ocr_data = process_ocr_data(table_text)

#     # Convert the DataFrame to a JSON response
#     json_response = ocr_data.to_json(orient='records')

#     return JsonResponse(json_response, safe=False)

# def process_ocr_data(ocr_text):
#     # Process the OCR text into lines
#     lines = ocr_text.split('\n')

#     # Extract potential column names from the header line
#     columns = extract_columns(lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line using spaces as separators
#         row = re.split(r'\s+', line.strip())

#         # Check if the row has the same number of elements as columns
#         if len(row) == len(columns):
#             # Create a dictionary with flexible keys based on the columns
#             data.append(dict(zip(columns, row)))

#     return pd.DataFrame(data)

# def extract_columns(header_line):
#     # Extract potential column names based on the structure of the header line
#     # Use spaces as separators
#     columns = re.split(r'\s+', header_line.strip())

#     # Filter out any empty strings or unwanted characters
#     columns = [col for col in columns if col]

#     return columns


# def read_word(file_path):
#     doc = docx.Document(file_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + "\n"
#     return text

# def read_word(uploaded_file):
#     doc = docx.Document(uploaded_file)

#     # Extract text from each paragraph
#     text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

#     # Split text into lines
#     lines = text.split('\n')

#     # Remove empty lines
#     lines = [line.strip() for line in lines if line.strip()]

#     # Check if there are any lines in the document
#     if not lines:
#         return {"columns": [], "data": []}

#     # Assume the first line contains the column headers
#     columns = re.split(r'\s{2,}', lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line into values using any whitespace as delimiter
#         values = re.split(r'\s{2,}', line)
#         data.append(values)

#     return {"columns": columns, "data": data}


# def read_word(word_file):
#     if word_file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#         # Read the Word document
#         doc = docx.Document(word_file)

#         # Initialize extracted data
#         extracted_data = {'columns': [], 'data': []}

#         # Iterate through tables in the document
#         for table in doc.tables:
#             columns = [cell.text.strip() for cell in table.rows[0].cells]
#             data = []

#             # Iterate through rows in the table
#             for row in table.rows[1:]:
#                 row_data = [cell.text.strip() for cell in row.cells]
#                 data.append(dict(zip(columns, row_data)))

#             # Append columns to the overall list (if not already present)
#             extracted_data['columns'].extend(col for col in columns if col not in extracted_data['columns'])
#             extracted_data['data'].extend(data)

#         return extracted_data
#     elif word_file.content_type.startswith('image'):
#         # If the Word document is an image, perform OCR
#         ocr_text = perform_ocr(word_file)
#         # Process ocr_text as needed
#         return {'text': ocr_text}
#     else:
#         raise ValueError("Unsupported file format for Word.")


def read_word(word_file):
    if word_file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # Read the Word document
        doc = docx.Document(word_file)

        # Initialize extracted data
        extracted_data = []

        # Iterate through tables in the document
        for table in doc.tables:
            columns = [cell.text.strip() for cell in table.rows[0].cells]
            data = []

            # Iterate through rows in the table
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(dict(zip(columns, row_data)))

            # Append data to the overall list
            extracted_data.extend(data)

        return extracted_data
    elif word_file.content_type.startswith('image'):
        # If the Word document is an image, perform OCR
        ocr_text = perform_ocr(word_file)
        # Process ocr_text as needed
        return {'text': ocr_text}
    else:
        raise ValueError("Unsupported file format for Word.")
    


def extract_text_from_pdf(uploaded_file):
    text = ""
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    num_pages = len(pdf_reader.pages)
    for page_num in range(num_pages):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

# def read_pdf_with_ocr(uploaded_file):
#     text_from_pdf = ""
#     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#     num_pages = len(pdf_reader.pages)
#     for page_num in range(num_pages):
#         page = pdf_reader.pages[page_num]
#         text_from_pdf += page.extract_text()

#     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#     num_pages = len(pdf_reader.pages)
#     images = []

#     for page_num in range(num_pages):
#         page = pdf_reader.pages[page_num]
#         if '/XObject' in page['/Resources']:
#             xObject = page['/Resources']['/XObject'].get_object()
#             if xObject.get_object() and '/Image' in xObject.get_object():
#                 images.append(xObject.get_object()['/Image'])

#     text_from_images = ""
#     for image in images:
#         if image['/ColorSpace'] == '/DeviceRGB':
#             image_data = image.get_object()
#             img = Image.open(io.BytesIO(image_data)).convert("RGB")
#             text_from_images += ocr_image(img)

#     return {"columns": [], "data": text_from_pdf + text_from_images}

# def read_pdf_with_ocr(uploaded_file):
#     # Create a temporary PDF file
#     temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
#     temp_pdf_path = temp_pdf.name

#     try:
#         # Write the content of the uploaded PDF file to the temporary file
#         with open(temp_pdf_path, 'wb') as temp_pdf_file:
#             temp_pdf_file.write(uploaded_file.read())

#         # Extract text using PyPDF2
#         text_from_pdf = ""
#         pdf_reader = PyPDF2.PdfReader(temp_pdf_path)
#         num_pages = len(pdf_reader.pages)

#         for page_num in range(num_pages):
#             page = pdf_reader.pages[page_num]
#             text_from_pdf += page.extract_text()

#         # Extract text from images using Tesseract OCR
#         images = []
#         for page_num in range(num_pages):
#             page = pdf_reader.pages[page_num]
#             if '/XObject' in page['/Resources']:
#                 xObject = page['/Resources']['/XObject'].getObject()
#                 if xObject and '/Image' in xObject:
#                     images.append(xObject['/Image'])

#         text_from_images = ""
#         for image in images:
#             image_data = image._data
#             img = Image.open(io.BytesIO(image_data)).convert("RGB")
#             text_from_images += pytesseract.image_to_string(img)

#         # Combine text from PDF and images
#         combined_text = text_from_pdf + text_from_images

#         # Split the combined text into lines
#         lines = combined_text.split('\n')

#         # Extract columns (assuming they are present in the first line)
#         columns = [col.strip() for col in lines[0].split()]

#         # Extract data rows
#         data = [row.split() for row in lines[1:] if row]

#         return {"columns": columns, "data": data}

#     except Exception as e:
#         print(f"Error processing PDF file: {e}")
#         return {"columns": [], "data": []}

#     finally:
#         # Close and remove the temporary PDF file
#         temp_pdf.close()
#         os.remove(temp_pdf_path)

# def read_pdf_with_ocr(uploaded_file):
#     text_from_pdf = ""

#     # Read the InMemoryUploadedFile content
#     content = uploaded_file.read()

#     # Create a PyMuPDF document
#     pdf_document = fitz.open("pdf", content)

#     for page_num in range(pdf_document.page_count):
#         page = pdf_document[page_num]
#         text_from_pdf += page.get_text()

#     # Split the text into lines
#     lines = text_from_pdf.split('\n')

#     # Remove empty lines
#     lines = [line.strip() for line in lines if line.strip()]

#     # Assume the first line contains the column headers
#     columns = re.split(r'\s{2,}', lines[0])

#     # Extract data from subsequent lines
#     data = []
#     for line in lines[1:]:
#         # Split the line into values using multiple spaces as delimiter
#         values = re.split(r'\s{2,}', line)
#         data.append(values)

#     return {"columns": columns, "data": data}


# def read_pdf_with_ocr(pdf_file):
#     # Convert PDF to images
#     images = convert_from_bytes(pdf_file.read())

#     # Perform OCR on each image
#     ocr_results = [image_to_string(image) for image in images]

#     # Extract data from OCR results
#     data = []
#     for ocr_result in ocr_results:
#         data.extend([line.split() for line in ocr_result.split('\n') if line])
    
#     # Create DataFrame
#     columns = data[0] if data else []
#     print("Columns:", columns)
#     print("Data:", data)
#     df = pd.DataFrame(data[1:], columns=None)
#     print(df)
 
#     return {"columns": columns, "data": df.to_dict(orient='records')}


# def read_pdf_with_ocr(uploaded_file):
#     images = convert_from_bytes(uploaded_file.read())

#     data = []
#     columns = None

#     for image in images:
#         ocr_result = pytesseract.image_to_string(image)
#         lines = [line.split() for line in ocr_result.split('\n') if line.strip()]

#         if not columns:
#             # Check if the first row contains any words
#             if lines and any(word.strip() for word in lines[0]):
#                 columns = lines[0]
#             else:
#                 # Use default column names if the first row is empty
#                 columns = [f"Column{i + 1}" for i in range(len(lines[0]))]

#         # Append subsequent lines to the existing data, assuming they are rows
#         for i, row in enumerate(lines[1:]):
#             if i < len(data):
#                 # Pad or truncate rows to match the number of columns
#                 data[i] += row[:len(columns)] + [''] * (len(columns) - len(row))
#             else:
#                 data.append(row[:len(columns)] + [''] * (len(columns) - len(row)))

#     # Create DataFrame
#     df = pd.DataFrame(data, columns=columns)

#     return {"columns": columns, "data": df.to_dict(orient='records')}

    
# def read_pdf_with_ocr(uploaded_file):
#     text_from_pdf = ""

#     # Extract text from PDF using PyPDF2
#     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#     num_pages = len(pdf_reader.pages)
#     for page_num in range(num_pages):
#         page = pdf_reader.pages[page_num]
#         text_from_pdf += page.extract_text()

#     # Extract text from images within the PDF using PyMuPDF (fitz) and Tesseract OCR
#     pdf_reader = PyPDF2.PdfReader(uploaded_file)
#     num_pages = len(pdf_reader.pages)
#     images_text = []

#     for page_num in range(num_pages):
#         page = pdf_reader.pages[page_num]
#         images = page.extract_images()
        
#         for img_index, image in enumerate(images):
#             image_index = image[0]
#             base_image = pdf_reader.images[image_index]
#             image_bytes = base_image.get_object()['/Bytes']
#             image_ext = base_image.get_object()['/Filter'][1:]

#             # Extract text from image using Tesseract OCR
#             if '/DCTDecode' in image_ext:
#                 image = Image.open(io.BytesIO(image_bytes))
#                 text_from_image = pytesseract.image_to_string(image)
#                 images_text.append(text_from_image)

#     # Combine text from PDF and images into a single string
#     all_text = text_from_pdf + "\n".join(images_text)

#     # Parse the text and create a list of dictionaries representing records
#     records = []
#     lines = all_text.split('\n')
#     if lines:
#         columns = [col.strip() for col in re.split(r'\s{2,}', lines[0].strip())]
#         for line in lines[1:]:
#             values = [val.strip() for val in re.split(r'\s{2,}', line.strip())]
#             record = dict(zip(columns, values))
#             records.append(record)

#     return {"columns": columns, "data": records}



# def read_pdf_with_ocr(pdf_file):
#     # Step 1: Convert PDF to an image
#     images = convert_pdf_to_images(pdf_file)

#     # Step 2: Perform OCR on each image
#     ocr_texts = [perform_ocr(image) for image in images]

#     # Step 3: Extract tabular data using tabula-py
#     extracted_data = {'columns': [], 'data': []}

#     for i, (image, ocr_text) in enumerate(zip(images, ocr_texts)):
#         tables = tabula.read_pdf(pdf_file, stream=True, pages="all") 

#         # Assume the first table on each page is the target table
#         if tables and tables[0].shape[0] > 0:
#             columns = list(tables[0].columns)
#             data = tables[0].to_dict(orient='records')

#             # Append columns to the overall list (if not already present)
#             extracted_data['columns'].extend(col for col in columns if col not in extracted_data['columns'])
#             extracted_data['data'].extend(data)

#     return extracted_data

# def read_pdf_with_ocr(pdf_file):
#     extracted_data = {'columns': [], 'data': []}
    
#     # Extract tables from all pages
#     tables = tabula.read_pdf(pdf_file, pages='all', multiple_tables=True, stream=True)

#     # Concatenate tables if they are split across multiple pages
#     for table in tables:
#         if not table.empty:
#             # Process each row in the table on the current page
#             columns = list(table.columns)
#             data = table.to_dict(orient='records')

#             # Append columns to the overall list (if not already present)
#             extracted_data['columns'].extend(col for col in columns if col not in extracted_data['columns'])
#             extracted_data['data'].extend(data) 

#     # Convert the extracted tables to a DataFrame or process as needed
#     df = pd.DataFrame(extracted_data['data'], columns=extracted_data['columns'])

#     # Convert DataFrame to a JSON-friendly format
#     extracted_data_json = df.to_json(orient='records')

#     return extracted_data_json


from tabula.io import read_pdf

def read_pdf_with_ocr(pdf_file):
    all_dfs = []  # List to store DataFrames from each page

    try:
        # Get the total number of pages in the PDF
        total_pages = len(tabula.read_pdf(pdf_file, pages='all', multiple_tables=True, stream=True))

        # Iterate through pages
        for page_num in range(1, total_pages + 1):
            # Extract tables from the current page
            tables = tabula.read_pdf(
                pdf_file,
                pages=page_num,
                multiple_tables=True,
                stream=True,
                pandas_options={'header': None}
            )

            # Process each table on the current page
            for table in tables:
                if not table.empty:
                    # Use the first row as header if it's the first page
                    if page_num == 1 and all_dfs == []:
                        column_names = table.iloc[0]
                        table = table[1:]

                    # Process each row in the table
                    data = table.to_dict(orient='records')

                    # Append data to the overall list only if it contains valid data
                    all_dfs.append(pd.DataFrame(data))

    except Exception as e:
        return json.dumps({"error": f"Error while extracting tables: {str(e)}"})

    # Concatenate all DataFrames into a single DataFrame
    print(all_dfs)
    final_df = pd.concat(all_dfs, ignore_index=True)
    print(final_df)

    # Rename columns with the first row of the first page
    if 'column_names' in locals():
        final_df.columns = column_names

    # Convert DataFrame to a JSON-friendly format
    extracted_data_json = final_df.to_json(orient='records')

    return extracted_data_json


# import logging

# logging.basicConfig(level=logging.ERROR)

# def read_pdf_with_ocr(uploaded_file):
#     if uploaded_file.content_type == 'application/pdf':
#         try:
#             # Use tabula to extract tables from the PDF
#             tables = tabula.read_pdf(uploaded_file, pages='all', multiple_tables=True)

#             # Check if any tables are extracted
#             if tables:
#                 # Convert the extracted tables to a JSON-friendly format
#                 extracted_data = {'tables': []}
#                 for i, table in enumerate(tables):
#                     table_data = table.to_dict(orient='records')
#                     extracted_data['tables'].append({'page': i + 1, 'data': table_data})

#                 # Convert the extracted_data to a JSON string
#                 extracted_data = json.dumps(extracted_data, ensure_ascii=False)

#                 # Convert JSON string to a Python object (if needed)
#                 extracted_data = json.loads(extracted_data)

#                 return extracted_data
#             else:
#                 logging.error("No tables extracted from the PDF.")
#                 return None
#         except Exception as e:
#             logging.error(f"Error extracting tables from PDF: {str(e)}")
#             return None
#     else:
#         logging.error("Unsupported file format for PDF.")
#         return None

    
def convert_pdf_to_images(pdf_file):
    # Use pdf2image to convert each page of the PDF to an image
    images = convert_from_bytes(pdf_file.read(), fmt='png')
    return images

def perform_ocr(image):
    # Use pytesseract to perform OCR on the image
    ocr_text = pytesseract.image_to_string(image)
    return ocr_text